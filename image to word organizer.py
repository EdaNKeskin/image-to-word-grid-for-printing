from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import os

# --- Settings ---
image_folder = "C:\\Users\\" # location of images
images_per_row = 3
image_width_in = 2.0  # inches in Word
target_dpi = 300      # for printing



# --- Create Word document ---
doc = Document()
row_images = []

#section = doc.sections[0] #landscape added
#section.orientation = WD_ORIENT.LANDSCAPE
#section.page_width, section.page_height = section.page_height, section.page_width


for filename in sorted(os.listdir(image_folder)):
    if filename.lower().endswith((".jpg", ".jpeg", ".png")):
        row_images.append(filename)

        if len(row_images) == images_per_row:
            table = doc.add_table(rows=1, cols=images_per_row)
            row_cells = table.rows[0].cells
            for i, img_name in enumerate(row_images):
                paragraph = row_cells[i].paragraphs[0]
                run = paragraph.add_run()

                # Open image and resize to print DPI (optional)
                img_path = os.path.join(image_folder, img_name)
                img = Image.open(img_path)
                img.thumbnail((int(image_width_in * target_dpi),
                               int(image_width_in * target_dpi)),
                               Image.Resampling.LANCZOS)
                temp_path = os.path.join(image_folder, f"temp_{img_name}")
                img.save(temp_path, quality=95)  # high JPEG quality

                run.add_picture(temp_path, width=Inches(image_width_in))
            row_images = []

if row_images:
    table = doc.add_table(rows=1, cols=len(row_images))
    row_cells = table.rows[0].cells
    for i, img_name in enumerate(row_images):
        paragraph = row_cells[i].paragraphs[0]
        run = paragraph.add_run()
        img_path = os.path.join(image_folder, img_name)
        img = Image.open(img_path)
        img.thumbnail((int(image_width_in * target_dpi),
                       int(image_width_in * target_dpi)),
                       Image.Resampling.LANCZOS)
        temp_path = os.path.join(image_folder, f"temp_{img_name}")
        img.save(temp_path, quality=95)
        run.add_picture(temp_path, width=Inches(image_width_in))

doc.save("all_pictures_grid.docx")
